import streamlit as st
import pdfplumber
import google.generativeai as genai
from io import BytesIO
from docx import Document

# ---------------------------
# PAGE CONFIG
# ---------------------------
st.set_page_config(page_title="AI PDF Summarizer Pro", page_icon="🤖", layout="wide")

# ---------------------------
# HEADER
# ---------------------------
st.markdown(
    """
    <div style="text-align:center; padding: 10px; background: linear-gradient(to right, #4facfe, #00f2fe); border-radius: 12px;">
        <h1 style="color:white;">📄 AI PDF Summarizer Pro</h1>
        <p style="color:white; font-size:18px;">Upload, Summarize, Ask Questions & Download Results</p>
    </div>
    """,
    unsafe_allow_html=True
)

# ---------------------------
# API CONFIG
# ---------------------------
api_key = "Your api key"

try:
    genai.configure(api_key=api_key)
    model = genai.GenerativeModel("gemini-2.5-flash")
    st.sidebar.success("✅ API Connected")
except Exception as e:
    st.sidebar.error(f"❌ API Error: {e}")
    model = None

# ---------------------------
# HELPERS
# ---------------------------
@st.cache_data
def read_pdf_content(uploaded_file):
    text = ""
    try:
        with pdfplumber.open(uploaded_file) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"
    except Exception as e:
        st.error(f"PDF reading error: {e}")
    return text

def chunk_text(text, chunk_size=1200, overlap=150):
    chunks = []
    start = 0
    while start < len(text):
        end = min(start + chunk_size, len(text))
        chunks.append(text[start:end])
        start = end - overlap if end - overlap > start else end
    return chunks

def summarize_text(chunks, model, style, language):
    summaries = []
    progress_bar = st.progress(0)
    status_text = st.empty()

    for i, chunk in enumerate(chunks):
        try:
            status_text.text(f"⏳ Summarizing chunk {i+1}/{len(chunks)}...")
            prompt = f"""
            Summarize the following text in {language} with {style} style:
            {chunk}
            """
            response = model.generate_content(prompt)
            summaries.append(getattr(response, "text", "[No summary]"))
        except Exception as e:
            st.error(f"Error summarizing chunk {i+1}: {e}")
            summaries.append(f"[Error in chunk {i+1}]")

        progress_bar.progress((i + 1) / len(chunks))

    status_text.empty()
    progress_bar.empty()
    return summaries

def refine_summary(summaries, model, language):
    try:
        prompt = f"Combine and refine these summaries into one clear summary in {language}:\n\n" + "\n\n".join(summaries)
        response = model.generate_content(prompt)
        return getattr(response, "text", "[No final summary]")
    except Exception as e:
        st.error(f"Error refining summary: {e}")
        return "[Refinement failed]"

def download_docx(summary_text, file_name):
    doc = Document()
    doc.add_heading("AI PDF Summary", level=1)
    doc.add_paragraph(summary_text)
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

# ---------------------------
# SIDEBAR SETTINGS
# ---------------------------
st.sidebar.header("⚙ Settings")
summary_style = st.sidebar.selectbox("📝 Summary Style", ["Concise", "Detailed", "Bullet Points"])
language = st.sidebar.selectbox("🌐 Language", ["English", "Hindi", "French"])
chunk_size = st.sidebar.slider("📏 Chunk Size", 500, 2000, 1200, 100)
overlap = st.sidebar.slider("🔁 Overlap", 0, 300, 150, 50)

# ---------------------------
# MAIN LAYOUT
# ---------------------------
col1, col2 = st.columns([2, 1])

with col1:
    uploaded_file = st.file_uploader("📤 Upload a PDF", type="pdf")

    if uploaded_file and model:
        st.success(f"✅ File uploaded: {uploaded_file.name}")
        
        with st.spinner("🔍 Extracting text..."):
            pdf_text = read_pdf_content(uploaded_file)

        if pdf_text.strip():
            st.info(f"📊 Extracted {len(pdf_text)} characters from PDF")

            with st.expander("📄 Preview Extracted Text", expanded=False):
                st.text_area("Text Preview", pdf_text[:1000] + "..." if len(pdf_text) > 1000 else pdf_text, height=200)

            if st.button("🚀 Generate Summary", use_container_width=True):
                text_chunks = chunk_text(pdf_text, chunk_size, overlap)
                summaries = summarize_text(text_chunks, model, summary_style, language)
                final_summary = refine_summary(summaries, model, language)

                st.session_state["summaries"] = summaries
                st.session_state["final_summary"] = final_summary

            # Always show summary if present
            if "summaries" in st.session_state and "final_summary" in st.session_state:
                st.markdown("### 📝 Chunk Summaries")
                for i, s in enumerate(st.session_state["summaries"]):
                    st.markdown(
                        f"""
                        <div style="background:#f1f1f1; padding:10px; border-radius:8px; margin-bottom:8px;">
                        <b>Chunk {i+1}:</b> {s}
                        </div>
                        """,
                        unsafe_allow_html=True
                    )

                st.markdown("### 📌 Final Refined Summary")
                st.markdown(
                    f"""
                    <div style="background:#e8f5e9; padding:15px; border-radius:10px; border:1px solid #c8e6c9;">
                    {st.session_state["final_summary"]}
                    </div>
                    """,
                    unsafe_allow_html=True
                )

                st.download_button(
                    "💾 Download TXT", 
                    data=st.session_state["final_summary"], 
                    file_name="summary.txt", 
                    mime="text/plain",
                    use_container_width=True
                )
                st.download_button(
                    "💾 Download DOCX", 
                    data=download_docx(st.session_state["final_summary"], "summary.docx"), 
                    file_name="summary.docx", 
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    use_container_width=True
                )

with col2:
    st.markdown("### ❓ Ask Questions from PDF")
    question = st.text_input("Enter your question")
    if st.button("🔍 Get Answer", use_container_width=True):
        try:
            prompt = f"Answer the question based only on this text:\n\n{pdf_text}\n\nQuestion: {question}"
            response = model.generate_content(prompt)
            st.markdown(
                f"""
                <div style="background:#fff3e0; padding:12px; border-radius:8px; border:1px solid #ffe0b2;">
                <b>💡 Answer:</b> {getattr(response, "text", "[No answer]")}
                </div>
                """,
                unsafe_allow_html=True
            )
        except Exception as e:
            st.error(f"Error answering question: {e}")


